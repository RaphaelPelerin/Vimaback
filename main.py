from fastapi import FastAPI, HTTPException, Depends, status, File, UploadFile, Form, Response, Body
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from pydantic import BaseModel, EmailStr, Field
from pymongo import MongoClient
from pymongo.server_api import ServerApi
from jose import JWTError, jwt
from passlib.context import CryptContext
from typing import Optional, List
from datetime import datetime, timedelta
import os
import certifi
from dotenv import load_dotenv
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
import smtplib
import docx
from docx.shared import Inches
import openai
import uuid
import base64
import io
from PIL import Image
import requests
import time
import random
from bson import ObjectId
import asyncio
import logging

# Load environment variables
load_dotenv()

# Setup logger
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# App initialization
app = FastAPI()

# Setup CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allows all origins
    allow_credentials=True,
    allow_methods=["*"],  # Allows all methods
    allow_headers=["*"],  # Allows all headers
)

# Create directories if they don't exist
os.makedirs("generated_images", exist_ok=True)
os.makedirs("generated_articles", exist_ok=True)

# Mount static directories
app.mount("/images", StaticFiles(directory="generated_images"), name="images")

# Global MongoDB client
mongodb_client = None
mongodb_database = None

def get_db():
    global mongodb_client, mongodb_database
    if mongodb_database is not None:
        return mongodb_database
        
    # Get MongoDB URI from environment variables
    MONGO_URI = os.getenv("MONGODB_URI")
    if not MONGO_URI:
        raise HTTPException(status_code=500, detail="MongoDB URI not configured")
    
    try:
        # Create a new client and connect to the server with SSL
        mongodb_client = MongoClient(
            MONGO_URI,
            tls=True,
            tlsAllowInvalidCertificates=True,  # Only for development
            serverSelectionTimeoutMS=5000,
            connectTimeoutMS=10000,
            retryWrites=True,
            w="majority"
        )
        
        # Send a ping to confirm a successful connection
        mongodb_client.admin.command('ping')
        
        # Store the database instance
        mongodb_database = mongodb_client.database
        return mongodb_database
        
    except Exception as e:
        print(f"Failed to connect to MongoDB: {e}")
        raise HTTPException(status_code=500, detail=f"Database connection failed: {str(e)}")

# Test database connection at startup
def test_database_connection():
    try:
        db = get_db()
        print("\n=== MongoDB Connection Test ===")
        print("✅ Successfully connected to MongoDB!")
        print("Database URL:", os.getenv("MONGODB_URI"))
        print("==============================\n")
        return True
    except Exception as e:
        print("\n=== MongoDB Connection Error ===")
        print("❌ Failed to connect to MongoDB!")
        print("Error details:", str(e))
        print("Database URL:", os.getenv("MONGODB_URI"))
        print("===============================\n")
        return False

@app.on_event("startup")
async def startup_event():
    print("\n=== Starting Server ===")
    if not test_database_connection():
        print("⚠️  Warning: Server starting with no database connection!")
    
    # Initialiser les prompts vidéo par défaut
    await initialize_default_video_prompts()
    
    print("=== Server Ready ===\n")

@app.on_event("shutdown")
async def shutdown_event():
    global mongodb_client
    if mongodb_client is not None:
        mongodb_client.close()
        print("MongoDB connection closed")

# Security utilities
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="token")

# Secret key and algorithm for JWT
SECRET_KEY = os.getenv("SECRET_KEY")
ALGORITHM = os.getenv("ALGORITHM")
ACCESS_TOKEN_EXPIRE_MINUTES = int(os.getenv("ACCESS_TOKEN_EXPIRE_MINUTES"))

# Models
class UserCreate(BaseModel):
    username: str
    email: EmailStr
    password: str
    role: str = "user"  # Default role is "user"

class UserInDB(BaseModel):
    id: Optional[str] = None
    _id: Optional[str] = None
    username: str
    email: EmailStr
    hashed_password: str
    role: str = "user"  # Default role is "user"
    api_key: Optional[str] = None  # Add API key field
    openai_api_key: Optional[str] = None  # Add OpenAI API key field
    
    class Config:
        orm_mode = True
        allow_population_by_field_name = True
        arbitrary_types_allowed = True

class Token(BaseModel):
    access_token: str
    token_type: str

class TokenData(BaseModel):
    email: Optional[str] = None

# Article Models
class ArticleCreate(BaseModel):
    title: str
    category: str
    summary: str
    content: str
    cover_image: Optional[str] = None

class ArticleResponse(BaseModel):
    id: str
    uid: str  # Adding unique identifier field
    title: str
    category: str
    summary: str
    content: dict  # Changed to dict to store structured content
    author: str
    created_at: datetime

class ArticleGenerate(BaseModel):
    theme: str
    email_to: EmailStr

# Video Models
class VideoCreate(BaseModel):
    title: str
    script: str
    style: str
    duration: int
    resolution: str
    voice_gender: str

class VideoResponse(BaseModel):
    id: str
    title: str
    status: str = "processing"
    video_url: Optional[str] = None
    thumbnail_url: Optional[str] = None
    author: str
    created_at: datetime

# Nouveaux modèles pour les prompts vidéo
class VideoPromptCreate(BaseModel):
    name: str
    description: str
    prompt_template: str
    category: str
    tags: List[str]

class VideoPromptResponse(BaseModel):
    id: str
    name: str
    description: str
    prompt_template: str
    category: str
    tags: List[str]
    created_by: str
    created_at: datetime

# Modèles pour le traitement CSV et la génération en lot
class VideoCSVEntry(BaseModel):
    content_type: str  # Nom du prompt
    theme: str

class BatchVideoCreate(BaseModel):
    entries: List[VideoCSVEntry]
    email_notification: Optional[EmailStr] = None

class BatchVideoStatus(BaseModel):
    id: str
    total_videos: int
    processed_videos: int
    status: str = "processing"  # processing, completed, failed
    start_time: datetime
    end_time: Optional[datetime] = None
    created_by: str

# Modèles pour l'intégration avec InVideo.ai
class InVideoTask(BaseModel):
    task_id: str
    status: str
    video_url: Optional[str] = None
    thumbnail_url: Optional[str] = None
    error_message: Optional[str] = None

# Modèles pour les statistiques
class VideoStatistics(BaseModel):
    total_videos: int
    completed_videos: int
    processing_videos: int
    failed_videos: int
    recent_activity: List[dict]

# Helper functions
def verify_password(plain_password, hashed_password):
    return pwd_context.verify(plain_password, hashed_password)

def get_password_hash(password):
    return pwd_context.hash(password)

def get_user(db, email: str):
    user = db.users.find_one({"email": email})
    if user:
        user["_id"] = str(user["_id"])  # Convert ObjectId to string
        return user
    return None

def authenticate_user(db, email: str, password: str):
    user = db.users.find_one({"email": email})
    if not user:
        return False
    if not verify_password(password, user["hashed_password"]):
        return False
    return UserInDB(**user)

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=15)
    # Add issue time
    to_encode.update({
        "exp": expire,
        "iat": datetime.utcnow()
    })
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

async def get_current_user(token: str = Depends(oauth2_scheme)):
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        email: str = payload.get("sub")
        if email is None:
            raise credentials_exception
        token_data = TokenData(email=email)
    except JWTError:
        raise credentials_exception
    
    db = get_db()
    user = get_user(db, email=token_data.email)
    if user is None:
        raise credentials_exception
    
    # Check if password has been changed since token was issued
    token_issued_at = datetime.fromtimestamp(payload.get("iat", 0))
    if "password_changed_at" in user and user["password_changed_at"] > token_issued_at:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Your password has been changed. Please login again",
            headers={"WWW-Authenticate": "Bearer"},
        )
    
    # Convert the user dictionary to a UserInDB model
    if isinstance(user, dict):
        # Copy the _id value to id if present
        if '_id' in user and 'id' not in user:
            user['id'] = str(user['_id'])
        # Create a UserInDB model
        user_model = UserInDB(**user)
        return user_model
    
    return user

async def get_current_admin_user(current_user: UserInDB = Depends(get_current_user)):
    if current_user.role not in ['admin', 'owner']:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="You don't have permission to access this resource"
        )
    return current_user

# Routes
@app.post("/register", status_code=status.HTTP_201_CREATED)
async def register_user(user: UserCreate):
    db = get_db()
    
    # Check if user already exists
    if db.users.find_one({"email": user.email}):
        raise HTTPException(
            status_code=400,
            detail="Email already registered"
        )
    
    # Create new user
    hashed_password = get_password_hash(user.password)
    user_dict = {
        "username": user.username,
        "email": user.email,
        "hashed_password": hashed_password,
        "role": "user",  # Always set new registrations to "user" role
        "created_at": datetime.utcnow()
    }
    
    result = db.users.insert_one(user_dict)
    if result.inserted_id:
        return {"message": "User created successfully"}
    else:
        raise HTTPException(
            status_code=500,
            detail="Failed to create user"
        )

@app.post("/token", response_model=Token)
async def login_for_access_token(form_data: OAuth2PasswordRequestForm = Depends()):
    db = get_db()
    user = authenticate_user(db, form_data.username, form_data.password)
    if not user:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect email or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    
    # Check if user has role field, if not add it
    user_data = db.users.find_one({"email": user.email})
    if "role" not in user_data:
        db.users.update_one(
            {"email": user.email},
            {"$set": {"role": "user"}}
        )
        user.role = "user"
    
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(
        data={"sub": user.email}, expires_delta=access_token_expires
    )
    
    return {"access_token": access_token, "token_type": "bearer"}

@app.get("/users/me")
async def read_users_me(current_user = Depends(get_current_user)):
    return {
        "username": current_user.username,
        "email": current_user.email,
        "role": current_user.role  # Include role in response
    }

# Article Endpoints
@app.post("/articles", response_model=ArticleResponse)
async def create_article(article: ArticleGenerate, current_user: UserInDB = Depends(get_current_user)):
    try:
        # Generate article content with default prompt type
        generated_content = generate_article_content(
            prompt_type="cybersecurity",
            theme=article.theme,
            additional_content=None
        )
        
        # Print the generated content for debugging
        print("\n=== Generated Content to Parse ===")
        print(generated_content)
        print("==============================\n")
        
        # Parse the structured response
        title = ""
        image_suggestion = ""
        article_content = ""
        hashtags = ""
        
        # Improved parsing logic
        if "===Titre===" in generated_content:
            title_parts = generated_content.split("===Titre===")
            if len(title_parts) > 1:
                next_section = "===IMAGE SUGGÉRÉE===" if "===IMAGE SUGGÉRÉE===" in title_parts[1] else (
                               "===ARTICLE===" if "===ARTICLE===" in title_parts[1] else 
                               "===HASHTAGS===" if "===HASHTAGS===" in title_parts[1] else "")
                
                if next_section:
                    title = title_parts[1].split(next_section)[0].strip()
                else:
                    title = title_parts[1].strip()
        
        if "===IMAGE SUGGÉRÉE===" in generated_content:
            img_parts = generated_content.split("===IMAGE SUGGÉRÉE===")
            if len(img_parts) > 1:
                next_section = "===ARTICLE===" if "===ARTICLE===" in img_parts[1] else (
                               "===HASHTAGS===" if "===HASHTAGS===" in img_parts[1] else "")
                
                if next_section:
                    image_suggestion = img_parts[1].split(next_section)[0].strip()
                else:
                    image_suggestion = img_parts[1].strip()
        
        if "===ARTICLE===" in generated_content:
            article_parts = generated_content.split("===ARTICLE===")
            if len(article_parts) > 1:
                next_section = "===HASHTAGS===" if "===HASHTAGS===" in article_parts[1] else ""
                
                if next_section:
                    article_content = article_parts[1].split(next_section)[0].strip()
                else:
                    article_content = article_parts[1].strip()
        
        if "===HASHTAGS===" in generated_content:
            hashtag_parts = generated_content.split("===HASHTAGS===")
            if len(hashtag_parts) > 1:
                hashtags = hashtag_parts[1].strip()
        
        # Print parsed content for debugging
        print("\n=== Parsed Content ===")
        print(f"Title: {title}")
        print(f"Image Suggestion: {image_suggestion}")
        print(f"Article Content: {article_content[:100]}...")  # Just print the beginning
        print(f"Hashtags: {hashtags}")
        print("==============================\n")
        
        # If no title was found in the content, use the theme
        if not title:
            title = f"Article sur {article.theme}"
        
        # Generate a unique identifier for this article
        article_uid = str(uuid.uuid4())
        
        # Generate the actual image using DALL-E if we have an image suggestion
        image_path = None
        image_error = None
        image_style = None


        user_record = db.users.find_one({"email": current_user.email})
        openai_api_key = user_record.get("openai_api_key") or os.getenv("OPENAI_API_KEY")
        print(f"OpenAI API key: {openai_api_key}")
        
        if image_suggestion:
            print(f"\n=== Génération d'image pour l'article {article_uid} ===")
            print(f"Suggestion d'image: {image_suggestion[:100]}...")
            
            # Attendre la génération de l'image
            image_path, image_error, image_style = await generate_image_from_prompt(image_suggestion, openai_api_key)
            
            if image_path:
                print(f"✅ Image générée avec succès: {image_path}")
            else:
                print(f"❌ Échec de la génération d'image: {image_error}")
        
        # Create structured content dictionary
        structured_content = {
            "title": title or f"Article sur {article.theme}",  # Fallback to theme if title is empty
            "image_suggestion": image_suggestion or "Aucune suggestion d'image fournie",  # Provide default if empty
            "article_content": article_content or "Contenu de l'article non disponible",  # Provide default if empty
            "hashtags": hashtags or "Aucun hashtag fourni",  # Provide default if empty
            "raw_content": generated_content,  # Keep the original content for reference
            "image_path": image_path,  # Path to the generated image if successful
            "image_style": image_style,  # Style used for the generated image
            "image_error": image_error  # Error message if image generation failed
        }
        
        # Create article document
        article_data = {
            "uid": article_uid,  # Add the unique identifier
            "title": structured_content["title"],  # Use the processed title
            "category": "cybersecurity",
            "summary": article.theme,
            "content": structured_content,  # Store the structured content
            "author": current_user.username,
            "user_email": current_user.email,  # Store user's email to link articles to users
            "created_at": datetime.utcnow()
        }
        
        # Save to database
        db = get_db()
        result = db.articles.insert_one(article_data)
        
        # Always save as Word document
        filename = f"Article_{article_uid}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = os.path.join("generated_articles", filename)
        os.makedirs("generated_articles", exist_ok=True)
        
        # Create Word document with structured content
        doc = docx.Document()
        doc.add_heading(structured_content["title"], 0)
        
        # Add identifier information
        doc.add_paragraph(f"Identifiant unique: {article_uid}")
        doc.add_paragraph(f"Généré par: {current_user.username}")
        doc.add_paragraph(f"Date: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        
        # Add a separator
        doc.add_paragraph("").paragraph_format.space_after = Inches(0.2)
        
        # If image was generated, add it to the document
        if image_path:
            doc.add_heading("Image Générée", level=2)
            doc.add_picture(image_path, width=Inches(6))
        
        # Add image suggestion
        if structured_content["image_suggestion"]:
            doc.add_heading("Image Suggérée", level=2)
            doc.add_paragraph(structured_content["image_suggestion"])
        
        # Add main content
        doc.add_paragraph(structured_content["article_content"])
        
        # Add hashtags
        if structured_content["hashtags"]:
            doc.add_heading("Hashtags", level=2)
            doc.add_paragraph(structured_content["hashtags"])
        
        doc.save(file_path)
        
        # Send email with structured content
        email_content = f"""
Bonjour,

Voici votre article généré sur le thème : {article.theme}

Identifiant unique de l'article : {article_uid}

=== TITRE ===
{structured_content["title"]}

=== IMAGE SUGGÉRÉE ===
{structured_content["image_suggestion"]}

=== CONTENU DE L'ARTICLE ===
{structured_content["article_content"]}

=== HASHTAGS ===
{structured_content["hashtags"]}

Cordialement,
Votre assistant de génération d'articles
"""
        
        # Attach generated image if available
        attachments = [file_path]
        if image_path:
            attachments.append(image_path)
        
        send_email(
            article.email_to,
            f"Article Généré : {structured_content['title']}",
            email_content,
            attachments
        )
        
        # Get the created article with its new ID
        created_article = db.articles.find_one({"_id": result.inserted_id})
        created_article["id"] = str(created_article.pop("_id"))
        
        return created_article
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/articles", response_model=List[ArticleResponse])
async def get_articles(current_user: UserInDB = Depends(get_current_user)):
    db = get_db()
    
    # Get all articles by the current user
    articles = []
    for article in db.articles.find({"user_email": current_user.email}):
        article["id"] = str(article.pop("_id"))
        articles.append(article)
    
    return articles

@app.get("/articles/{article_uid}", response_model=ArticleResponse)
async def get_article_by_uid(article_uid: str, current_user: UserInDB = Depends(get_current_user)):
    db = get_db()
    
    # Retrieve the article with the specified UID, only if it belongs to the current user
    article = db.articles.find_one({"uid": article_uid, "user_email": current_user.email})
    
    if not article:
        raise HTTPException(status_code=404, detail="Article not found or you don't have permission to view it")
    
    article["id"] = str(article.pop("_id"))
    return article

# Video Endpoints
@app.post("/videos", response_model=VideoResponse)
async def create_video(video: VideoCreate, current_user: UserInDB = Depends(get_current_user)):
    db = get_db()
    
    video_dict = video.dict()
    video_dict.update({
        "author": current_user.username,
        "created_at": datetime.utcnow(),
        "status": "processing"
    })
    
    result = db.videos.insert_one(video_dict)
    
    # Get the created video with its new ID
    created_video = db.videos.find_one({"_id": result.inserted_id})
    created_video["id"] = str(created_video.pop("_id"))
    
    # In a real application, you would start a background task to generate the video here
    
    return created_video

@app.get("/videos", response_model=List[VideoResponse])
async def get_videos(current_user: UserInDB = Depends(get_current_user)):
    db = get_db()
    
    # Get all videos by the current user
    videos = []
    for video in db.videos.find({"author": current_user.username}):
        video["id"] = str(video.pop("_id"))
        videos.append(video)
    
    return videos

# Endpoints de prompts vidéo
@app.post("/video-prompts", response_model=VideoPromptResponse)
async def create_video_prompt(prompt: VideoPromptCreate, current_user: UserInDB = Depends(get_current_user)):
    db = get_db()
    
    # Créer le document prompt
    prompt_dict = prompt.dict()
    prompt_dict.update({
        "created_by": current_user.username,
        "created_at": datetime.utcnow()
    })
    
    result = db.video_prompts.insert_one(prompt_dict)
    
    # Récupérer le prompt créé avec son nouvel ID
    created_prompt = db.video_prompts.find_one({"_id": result.inserted_id})
    created_prompt["id"] = str(created_prompt.pop("_id"))
    
    return created_prompt

@app.get("/video-prompts", response_model=List[VideoPromptResponse])
async def get_video_prompts(current_user: UserInDB = Depends(get_current_user)):
    db = get_db()
    
    # Récupérer tous les prompts (filtré par utilisateur ou tous les prompts)
    prompts = []
    # On affiche à la fois les prompts de l'utilisateur et les prompts système
    query = {"$or": [{"created_by": current_user.username}, {"is_system": True}]}
    
    for prompt in db.video_prompts.find(query):
        prompt["id"] = str(prompt.pop("_id"))
        prompts.append(prompt)
    
    return prompts

@app.put("/video-prompts/{prompt_id}", response_model=VideoPromptResponse)
async def update_video_prompt(prompt_id: str, prompt_data: VideoPromptCreate, current_user: UserInDB = Depends(get_current_user)):
    db = get_db()
    
    # Vérifier si le prompt existe et appartient à l'utilisateur actuel
    existing_prompt = db.video_prompts.find_one({"_id": prompt_id, "created_by": current_user.username})
    
    if not existing_prompt:
        raise HTTPException(status_code=404, detail="Prompt non trouvé ou vous n'avez pas la permission de le modifier")
    
    # Mettre à jour le prompt
    updated_data = prompt_data.dict()
    db.video_prompts.update_one(
        {"_id": prompt_id},
        {"$set": updated_data}
    )
    
    # Récupérer le prompt mis à jour
    updated_prompt = db.video_prompts.find_one({"_id": prompt_id})
    updated_prompt["id"] = str(updated_prompt.pop("_id"))
    
    return updated_prompt

@app.delete("/video-prompts/{prompt_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_video_prompt(prompt_id: str, current_user: UserInDB = Depends(get_current_user)):
    db = get_db()
    
    # Vérifier si le prompt existe et appartient à l'utilisateur actuel
    existing_prompt = db.video_prompts.find_one({"_id": prompt_id, "created_by": current_user.username})
    
    if not existing_prompt:
        raise HTTPException(status_code=404, detail="Prompt non trouvé ou vous n'avez pas la permission de le supprimer")
    
    # Supprimer le prompt
    db.video_prompts.delete_one({"_id": prompt_id})
    
    return None

# Endpoints pour le traitement CSV et la génération en lot
@app.post("/videos/batch", response_model=BatchVideoStatus)
async def create_batch_videos(batch_request: BatchVideoCreate, current_user: UserInDB = Depends(get_current_user)):
    db = get_db()
    
    # Créer un document pour suivre le statut du lot
    batch_status = {
        "total_videos": len(batch_request.entries),
        "processed_videos": 0,
        "status": "processing",
        "start_time": datetime.utcnow(),
        "created_by": current_user.username,
        "entries": batch_request.dict().get("entries", []),
        "email_notification": batch_request.email_notification
    }
    
    result = db.batch_videos.insert_one(batch_status)
    batch_id = str(result.inserted_id)
    
    # Dans une application réelle, vous lanceriez ici une tâche en arrière-plan pour traiter le lot
    # Pour l'exemple, nous allons simplement créer les entrées de vidéos
    
    # Mettre à jour la réponse avec l'ID du lot
    created_batch = db.batch_videos.find_one({"_id": result.inserted_id})
    created_batch["id"] = str(created_batch.pop("_id"))
    
    # Lancer le traitement des vidéos en arrière-plan
    # Ici, vous pourriez utiliser Celery, FastAPI BackgroundTasks, etc.
    # Nous simulerons simplement la création des entrées vidéo
    
    return created_batch

@app.get("/videos/batch/{batch_id}", response_model=BatchVideoStatus)
async def get_batch_status(batch_id: str, current_user: UserInDB = Depends(get_current_user)):
    db = get_db()
    
    try:
        # Récupérer le statut du lot
        batch = db.batch_videos.find_one({
            "_id": ObjectId(batch_id), 
            "created_by": current_user.username
        })
        
        if not batch:
            raise HTTPException(status_code=404, detail="Lot non trouvé ou vous n'avez pas la permission de le consulter")
        
        batch["id"] = str(batch.pop("_id"))
        
        return batch
    except Exception as e:
        print(f"Erreur lors de la récupération du batch {batch_id}: {str(e)}")
        raise HTTPException(status_code=404, detail=f"Lot non trouvé: {str(e)}")

# Endpoint pour les statistiques
@app.get("/stats/videos", response_model=VideoStatistics)
async def get_video_statistics(current_user: UserInDB = Depends(get_current_user)):
    db = get_db()
    
    # Récupérer le nombre total de vidéos
    total_videos = db.videos.count_documents({"author": current_user.username})
    
    # Récupérer le nombre de vidéos par statut
    completed_videos = db.videos.count_documents({"author": current_user.username, "status": "completed"})
    processing_videos = db.videos.count_documents({"author": current_user.username, "status": "processing"})
    failed_videos = db.videos.count_documents({"author": current_user.username, "status": "failed"})
    
    # Récupérer les 5 dernières activités
    recent_videos = list(db.videos.find(
        {"author": current_user.username},
        sort=[("created_at", -1)],
        limit=5
    ))
    
    # Formater les activités récentes
    recent_activity = []
    for video in recent_videos:
        video["id"] = str(video.pop("_id"))
        recent_activity.append({
            "id": video["id"],
            "title": video["title"],
            "status": video["status"],
            "created_at": video["created_at"]
        })
    
    return {
        "total_videos": total_videos,
        "completed_videos": completed_videos,
        "processing_videos": processing_videos,
        "failed_videos": failed_videos,
        "recent_activity": recent_activity
    }

# Health check endpoint
@app.get("/health")
def health_check():
    return {"status": "ok"}

# Endpoint pour générer un scénario vidéo
@app.post("/videos/generate-script")
async def generate_video_script_endpoint(
    data: dict,
    current_user: UserInDB = Depends(get_current_user)
):
    if not data.get("prompt_id") and not data.get("prompt_template"):
        raise HTTPException(
            status_code=400, 
            detail="Vous devez fournir soit un ID de prompt, soit un template de prompt"
        )
    
    if not data.get("theme"):
        raise HTTPException(status_code=400, detail="Vous devez fournir un thème pour la vidéo")
    
    db = get_db()
    
    # Récupérer le template de prompt
    prompt_template = None
    if data.get("prompt_id"):
        prompt = db.video_prompts.find_one({
            "_id": data.get("prompt_id"),
            "$or": [{"created_by": current_user.username}, {"is_system": True}]
        })
        
        if not prompt:
            raise HTTPException(status_code=404, detail="Prompt non trouvé ou vous n'avez pas la permission de l'utiliser")
        
        prompt_template = prompt.get("prompt_template")
    else:
        prompt_template = data.get("prompt_template")
    
    # Générer le scénario
    try:
        script = generate_video_script(prompt_template, data.get("theme"))
        
        # Enregistrer le scénario dans la base de données
        script_entry = {
            "theme": data.get("theme"),
            "prompt_template": prompt_template,
            "generated_script": script,
            "author": current_user.username,
            "created_at": datetime.utcnow()
        }
        
        db.video_scripts.insert_one(script_entry)
        
        return {
            "script": script,
            "theme": data.get("theme")
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erreur lors de la génération du scénario: {str(e)}")

# Fonction pour générer un scénario vidéo
def generate_video_script(prompt_template: str, theme: str):
    # Set OpenAI API key
    openai.api_key = os.getenv("OPENAI_API_KEY")
    
    # Remplacer les variables dans le template
    full_prompt = prompt_template.replace("[THEME]", theme)
    
    # Ajouter des instructions pour le format du scénario
    full_prompt += """

Format du scénario:
- Introduction attrayante
- Points clés à aborder
- Narration claire et concise
- Conclusion avec appel à l'action
- Durée adaptée à une vidéo de 2-3 minutes
"""
    
    # Imprimer le prompt envoyé à OpenAI
    print("\n=== Prompt envoyé à OpenAI pour scénario vidéo ===")
    print(full_prompt)
    print("===============================================\n")
    
    # Générer le contenu avec OpenAI
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "Vous êtes un scénariste professionnel spécialisé dans la création de scripts vidéo engageants et informatifs."},
            {"role": "user", "content": full_prompt}
        ],
        temperature=0.7,
        max_tokens=1500
    )
    
    # Récupérer le contenu généré
    script = response.choices[0].message['content']
    
    # Imprimer la réponse de ChatGPT
    print("\n=== Réponse de ChatGPT pour scénario vidéo ===")
    print(script)
    print("========================================\n")
    
    return script

# Classe pour gérer l'intégration avec InVideo.ai
class InVideoClient:
    def __init__(self, api_key=None):
        self.api_key = api_key or os.getenv("INVIDEO_API_KEY")
        self.base_url = "https://api.invideo.io"
    
    def create_video(self, script, title, style="professional", duration=60, voice="female"):
        """
        Crée une vidéo via l'API InVideo.ai
        
        Cette méthode effectue une requête à l'API InVideo.ai pour générer une vidéo
        à partir du script fourni.
        """
        print(f"\n=== Soumission d'une vidéo à InVideo.ai ===")
        print(f"Titre: {title}")
        print(f"Style: {style}")
        print(f"Durée: {duration} secondes")
        print(f"Voix: {voice}")
        print(f"Script (extrait): {script[:100]}...")
        
        try:
            # Préparation des données pour l'API
            payload = {
                "title": title,
                "script": script,
                "style": style,
                "duration": duration,
                "voice": voice,
                "language": "fr" # Langue française pour le script
            }
            
            # En-têtes pour l'API
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {self.api_key}"
            }
            
            # Appel à l'API InVideo
            response = requests.post(
                f"{self.base_url}/v1/videos",
                json=payload,
                headers=headers
            )
            
            # Vérifier si la requête a réussi
            if response.status_code in [200, 201, 202]:
                data = response.json()
                task_id = data.get("id") or data.get("task_id") or str(uuid.uuid4())
                
                print(f"✅ Vidéo soumise avec succès. ID de tâche: {task_id}")
                
                return {
                    "task_id": task_id,
                    "status": "processing"
                }
            else:
                # En cas d'erreur, on renvoie les détails
                error_message = f"Erreur {response.status_code} lors de la création de la vidéo: {response.text}"
                print(f"❌ {error_message}")
                
                # Retourner un ID temporaire et le statut failed
                return {
                    "task_id": str(uuid.uuid4()),
                    "status": "failed",
                    "error_message": error_message
                }
                
        except Exception as e:
            error_message = f"Exception lors de la création de la vidéo: {str(e)}"
            print(f"❌ {error_message}")
            
            # Créer un ID temporaire en cas d'échec
            return {
                "task_id": str(uuid.uuid4()),
                "status": "failed",
                "error_message": error_message
            }
    
    def get_video_status(self, task_id):
        """
        Récupère le statut d'une vidéo via l'API InVideo.ai
        """
        try:
            # En-têtes pour l'API
            headers = {
                "Authorization": f"Bearer {self.api_key}"
            }
            
            # Appel à l'API pour vérifier le statut
            response = requests.get(
                f"{self.base_url}/v1/videos/{task_id}",
                headers=headers
            )
            
            # Vérifier si la requête a réussi
            if response.status_code == 200:
                data = response.json()
                
                # Mapper le statut retourné par l'API au format interne
                api_status = data.get("status", "").lower()
                
                if api_status in ["completed", "done", "ready"]:
                    status = "completed"
                    video_url = data.get("video_url") or f"https://api.invideo.io/videos/{task_id}/download"
                    thumbnail_url = data.get("thumbnail_url") or f"https://api.invideo.io/videos/{task_id}/thumbnail"
                    
                    return {
                        "task_id": task_id,
                        "status": status,
                        "video_url": video_url,
                        "thumbnail_url": thumbnail_url
                    }
                    
                elif api_status in ["failed", "error"]:
                    return {
                        "task_id": task_id,
                        "status": "failed",
                        "error_message": data.get("error") or "La génération de la vidéo a échoué"
                    }
                    
                else:
                    # Statut par défaut: en cours de traitement
                    return {
                        "task_id": task_id,
                        "status": "processing"
                    }
            else:
                # En cas d'erreur lors de la vérification du statut
                return {
                    "task_id": task_id,
                    "status": "failed",
                    "error_message": f"Erreur {response.status_code} lors de la vérification du statut: {response.text}"
                }
                
        except Exception as e:
            # En cas d'exception lors de la vérification
            return {
                "task_id": task_id,
                "status": "failed",
                "error_message": f"Exception lors de la vérification du statut: {str(e)}"
            }

# Endpoint pour traiter un lot de vidéos
@app.post("/videos/batch/{batch_id}/process")
async def process_batch_videos(
    batch_id: str,
    current_user: UserInDB = Depends(get_current_user)
):
    db = get_db()
    
    try:
        # Vérifier que le lot existe et appartient à l'utilisateur
        batch = db.batch_videos.find_one({
            "_id": ObjectId(batch_id),
            "created_by": current_user.username
        })
        
        if not batch:
            raise HTTPException(status_code=404, detail="Lot non trouvé ou vous n'avez pas la permission de le modifier")
        
        # Vérifier le statut actuel
        if batch.get("status") not in ["processing", "not_started"]:
            raise HTTPException(status_code=400, detail=f"Le lot est déjà dans l'état: {batch.get('status')}")
        
        # Mettre à jour le statut du lot
        db.batch_videos.update_one(
            {"_id": ObjectId(batch_id)},
            {"$set": {
                "status": "processing",
                "start_processing_time": datetime.utcnow()
            }}
        )
        
        # Récupérer tous les scénarios pour ce lot
        scripts = list(db.batch_scripts.find({"batch_id": batch_id}))
        
        if not scripts:
            raise HTTPException(status_code=404, detail="Aucun scénario trouvé pour ce lot")
        
        # Initialiser des variables pour le suivi
        processed = 0
        total = len(scripts)
        
        # Créer une instance du client InVideo avec la clé API
        invideo_client = InVideoClient()
        
        # Traiter chaque scénario et créer une vidéo
        for script in scripts:
            try:
                # Vérifier si le scénario a déjà été traité
                if script.get("status") == "video_generated":
                    processed += 1
                    continue
                
                # Vérifier si le script est valide
                if not script.get("script"):
                    print(f"⚠️ Scénario {script.get('_id')} sans contenu - ignoré")
                    continue
                
                # Préparer les détails de la vidéo
                title = f"Vidéo sur {script.get('theme')}"
                script_content = script.get("script")
                
                # Créer la vidéo via l'API InVideo
                video_task = invideo_client.create_video(
                    script=script_content,
                    title=title,
                    style="professional",  # style par défaut
                    duration=90,  # durée par défaut en secondes
                    voice="female"  # voix par défaut
                )
                
                # Créer une entrée vidéo dans la base de données
                video_data = {
                    "title": title,
                    "script": script_content,
                    "style": "professional",
                    "duration": 90,
                    "resolution": "1080p", 
                    "voice_gender": "female",
                    "status": video_task.get("status", "processing"),
                    "author": current_user.username,
                    "created_at": datetime.utcnow(),
                    "batch_id": batch_id,
                    "script_id": str(script.get("_id")),
                    "task_id": video_task.get("task_id")
                }
                
                # Si la tâche a échoué immédiatement
                if video_task.get("status") == "failed":
                    video_data["error_message"] = video_task.get("error_message")
                
                # Sauvegarder dans la base de données
                result = db.videos.insert_one(video_data)
                video_id = str(result.inserted_id)
                
                # Mettre à jour le scénario avec le statut
                db.batch_scripts.update_one(
                    {"_id": script.get("_id")},
                    {"$set": {
                        "status": "video_processing",
                        "video_id": video_id
                    }}
                )
                
                processed += 1
                
                # Mettre à jour le statut du lot
                db.batch_videos.update_one(
                    {"_id": ObjectId(batch_id)},
                    {"$set": {
                        "processed_videos": processed
                    }}
                )
                
                # Attendre un peu pour ne pas surcharger l'API
                time.sleep(1)
                
            except Exception as e:
                print(f"Erreur lors du traitement du script {script.get('_id')}: {str(e)}")
                # Continuer avec le script suivant même en cas d'erreur
        
        # Mettre à jour le statut final du lot
        final_status = "processing"  # Reste en "processing" car les vidéos sont en cours de génération
        db.batch_videos.update_one(
            {"_id": ObjectId(batch_id)},
            {"$set": {
                "status": final_status,
                "processed_videos": processed,
                "last_updated_at": datetime.utcnow()
            }}
        )
        
        # Lancer en arrière-plan un processus pour vérifier régulièrement le statut des vidéos
        # (Dans une application réelle, cela serait fait par un scheduler ou un worker)
        
        return {
            "batch_id": batch_id,
            "total_videos": total,
            "processed_videos": processed,
            "status": final_status,
            "message": "Vidéos en cours de génération. La durée du processus dépend du nombre de vidéos."
        }
    except Exception as e:
        print(f"Erreur lors du traitement du lot {batch_id}: {str(e)}")
        # En cas d'erreur, mettre à jour le statut du lot
        db.batch_videos.update_one(
            {"_id": ObjectId(batch_id)},
            {"$set": {
                "status": "failed",
                "error_message": str(e),
                "end_time": datetime.utcnow()
            }}
        )
        raise HTTPException(status_code=500, detail=f"Erreur lors du traitement du lot: {str(e)}")

# Endpoint pour vérifier et mettre à jour le statut des vidéos
@app.post("/videos/check-status")
async def check_videos_status(current_user: UserInDB = Depends(get_current_user)):
    db = get_db()
    
    # Récupérer toutes les vidéos en cours de traitement
    processing_videos = list(db.videos.find({
        "author": current_user.username,
        "status": "processing"
    }))
    
    if not processing_videos:
        return {
            "message": "Aucune vidéo en cours de traitement",
            "updated_count": 0
        }
    
    # Créer une instance du client InVideo
    invideo_client = InVideoClient()
    
    # Vérifier le statut de chaque vidéo
    updated_count = 0
    
    for video in processing_videos:
        try:
            # Récupérer le statut de la vidéo
            task_id = video.get("task_id")
            if not task_id:
                continue
                
            video_status = invideo_client.get_video_status(task_id)
            
            # Mettre à jour le statut de la vidéo si nécessaire
            if video_status.get("status") != "processing":
                update_data = {
                    "status": video_status.get("status")
                }
                
                # Ajouter les URLs si la vidéo est terminée
                if video_status.get("status") == "completed":
                    update_data["video_url"] = video_status.get("video_url")
                    update_data["thumbnail_url"] = video_status.get("thumbnail_url")
                
                # Ajouter le message d'erreur si la vidéo a échoué
                if video_status.get("status") == "failed":
                    update_data["error_message"] = video_status.get("error_message")
                
                # Mettre à jour la vidéo
                db.videos.update_one(
                    {"_id": video.get("_id")},
                    {"$set": update_data}
                )
                
                updated_count += 1
        except Exception as e:
            print(f"Erreur lors de la vérification du statut de la vidéo {video.get('_id')}: {str(e)}")
    
    return {
        "message": f"{updated_count} vidéos ont été mises à jour",
        "updated_count": updated_count
    }

# Endpoint pour récupérer une vidéo par son ID
@app.get("/videos/{video_id}", response_model=VideoResponse)
async def get_video_by_id(video_id: str, current_user: UserInDB = Depends(get_current_user)):
    db = get_db()
    
    # Récupérer la vidéo
    video = db.videos.find_one({"_id": video_id, "author": current_user.username})
    
    if not video:
        raise HTTPException(status_code=404, detail="Vidéo non trouvée ou vous n'avez pas la permission de la consulter")
    
    video["id"] = str(video.pop("_id"))
    
    return video

# Endpoint pour télécharger un fichier CSV
@app.post("/upload/csv")
async def upload_csv_file(
    file: UploadFile = File(...),
    generate_scripts: bool = Form(False),
    current_user: UserInDB = Depends(get_current_user)
):
    if not file.filename.endswith('.csv'):
        raise HTTPException(status_code=400, detail="Le fichier doit être au format CSV")
    
    try:
        # Lire le contenu du fichier
        contents = await file.read()
        contents_str = contents.decode('utf-8')
        
        # Parser le CSV
        import csv
        from io import StringIO
        
        csv_file = StringIO(contents_str)
        csv_reader = csv.reader(csv_file)
        
        # Lire les en-têtes
        headers = next(csv_reader)
        
        # Vérifier si les en-têtes requis sont présents
        if len(headers) < 2:
            raise HTTPException(
                status_code=400, 
                detail="Le fichier CSV doit contenir au moins 2 colonnes : 'Type de contenu' et 'Thème'"
            )
        
        # Traiter les lignes
        entries = []
        for row in csv_reader:
            if len(row) >= 2 and row[0] and row[1]:  # S'assurer que les colonnes requises ne sont pas vides
                entries.append({
                    "content_type": row[0],
                    "theme": row[1]
                })
        
        if not entries:
            raise HTTPException(status_code=400, detail="Aucune entrée valide trouvée dans le fichier CSV")
        
        # Créer un lot de vidéos
        batch_status = {
            "total_videos": len(entries),
            "processed_videos": 0,
            "status": "processing",
            "start_time": datetime.utcnow(),
            "created_by": current_user.username,
            "entries": entries,
            "filename": file.filename
        }
        
        db = get_db()
        result = db.batch_videos.insert_one(batch_status)
        batch_id = str(result.inserted_id)
        
        # Si la génération de scripts est demandée, lancer le processus
        if generate_scripts:
            # Créer des entrées pour les scénarios
            scripts = []
            
            for i, entry in enumerate(entries):
                try:
                    # Trouver le prompt correspondant
                    prompt = db.video_prompts.find_one({"name": entry["content_type"]})
                    
                    if not prompt:
                        # Si le prompt n'existe pas, créer une entrée d'erreur
                        scripts.append({
                            "batch_id": batch_id,
                            "content_type": entry["content_type"],
                            "theme": entry["theme"],
                            "error": f"Prompt non trouvé: {entry['content_type']}",
                            "status": "error",
                            "created_at": datetime.utcnow(),
                            "author": current_user.username,
                            "order": i
                        })
                        continue
                    
                    # Générer le scénario
                    script_content = generate_video_script(
                        prompt_template=prompt["prompt_template"],
                        theme=entry["theme"]
                    )
                    
                    # Enregistrer le scénario
                    script_entry = {
                        "batch_id": batch_id,
                        "content_type": entry["content_type"],
                        "theme": entry["theme"],
                        "script": script_content,
                        "status": "generated",
                        "created_at": datetime.utcnow(),
                        "author": current_user.username,
                        "order": i
                    }
                    
                    scripts.append(script_entry)
                    
                except Exception as e:
                    # En cas d'erreur, enregistrer l'erreur
                    scripts.append({
                        "batch_id": batch_id,
                        "content_type": entry["content_type"],
                        "theme": entry["theme"],
                        "error": str(e),
                        "status": "error",
                        "created_at": datetime.utcnow(),
                        "author": current_user.username,
                        "order": i
                    })
            
            # Insérer tous les scénarios
            if scripts:
                db.batch_scripts.insert_many(scripts)
            
            # Mettre à jour le statut du lot
            db.batch_videos.update_one(
                {"_id": result.inserted_id},
                {"$set": {"scripts_generated": True}}
            )
        
        return {
            "batch_id": batch_id,
            "filename": file.filename,
            "entries_count": len(entries),
            "status": "processing",
            "scripts_generated": generate_scripts
        }
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erreur lors du traitement du fichier CSV: {str(e)}")

# Endpoint pour récupérer les scénarios d'un lot
@app.get("/videos/batch/{batch_id}/scripts")
async def get_batch_scripts(
    batch_id: str,
    current_user: UserInDB = Depends(get_current_user)
):
    db = get_db()
    
    try:
        # Vérifier que le lot existe et appartient à l'utilisateur
        batch = db.batch_videos.find_one({
            "_id": ObjectId(batch_id),
            "created_by": current_user.username
        })
        
        if not batch:
            raise HTTPException(status_code=404, detail="Lot non trouvé ou vous n'avez pas la permission de le consulter")
        
        # Récupérer les scénarios
        scripts = list(db.batch_scripts.find(
            {"batch_id": batch_id},
            sort=[("order", 1)]  # Trier par ordre d'origine
        ))
        
        # Formater les scripts pour la réponse
        formatted_scripts = []
        for script in scripts:
            formatted_script = {
                "id": str(script["_id"]),
                "content_type": script.get("content_type", ""),
                "theme": script.get("theme", ""),
                "script": script.get("script", ""),
                "status": script.get("status", ""),
                "error": script.get("error", None),
                "created_at": script.get("created_at", datetime.utcnow())
            }
            formatted_scripts.append(formatted_script)
        
        return {
            "batch_id": batch_id,
            "scripts": formatted_scripts
        }
    except Exception as e:
        print(f"Erreur lors de la récupération des scripts du batch {batch_id}: {str(e)}")
        raise HTTPException(status_code=404, detail=f"Lot non trouvé: {str(e)}")

# Endpoint pour mettre à jour un scénario
@app.put("/videos/batch/{batch_id}/scripts/{script_id}")
async def update_batch_script(
    batch_id: str,
    script_id: str,
    script_data: dict,
    current_user: UserInDB = Depends(get_current_user)
):
    db = get_db()
    
    try:
        # Vérifier que le lot existe et appartient à l'utilisateur
        batch = db.batch_videos.find_one({
            "_id": ObjectId(batch_id),
            "created_by": current_user.username
        })
        
        if not batch:
            raise HTTPException(status_code=404, detail="Lot non trouvé ou vous n'avez pas la permission de le modifier")
        
        # Mettre à jour le scénario
        result = db.batch_scripts.update_one(
            {"_id": ObjectId(script_id), "batch_id": batch_id},
            {"$set": {
                "script": script_data.get("script"),
                "edited": True,
                "updated_at": datetime.utcnow()
            }}
        )
        
        if result.modified_count == 0:
            raise HTTPException(status_code=404, detail="Scénario non trouvé ou aucune modification apportée")
        
        # Récupérer le scénario mis à jour
        updated_script = db.batch_scripts.find_one({"_id": ObjectId(script_id)})
        
        if not updated_script:
            raise HTTPException(status_code=404, detail="Scénario non trouvé après mise à jour")
        
        return {
            "id": str(updated_script["_id"]),
            "content_type": updated_script.get("content_type", ""),
            "theme": updated_script.get("theme", ""),
            "script": updated_script.get("script", ""),
            "status": updated_script.get("status", ""),
            "edited": updated_script.get("edited", False),
            "updated_at": updated_script.get("updated_at", datetime.utcnow())
        }
    except Exception as e:
        print(f"Erreur lors de la mise à jour du script {script_id} dans le batch {batch_id}: {str(e)}")
        raise HTTPException(status_code=404, detail=f"Erreur lors de la mise à jour: {str(e)}")

# Endpoint pour vérifier le statut d'une vidéo en cours de génération
@app.get("/videos/{video_id}/status")
async def check_video_status(
    video_id: str,
    current_user: UserInDB = Depends(get_current_user)
):
    db = get_db()
    
    try:
        # Récupérer la vidéo
        video = db.videos.find_one({
            "_id": ObjectId(video_id),
            "author": current_user.username
        })
        
        if not video:
            raise HTTPException(status_code=404, detail="Vidéo non trouvée ou vous n'avez pas la permission de la consulter")
        
        # Si la vidéo est déjà complétée ou a échoué, renvoyer simplement son statut
        if video.get("status") in ["completed", "failed"]:
            return {
                "id": str(video["_id"]),
                "title": video.get("title"),
                "status": video.get("status"),
                "video_url": video.get("video_url"),
                "thumbnail_url": video.get("thumbnail_url"),
                "error_message": video.get("error_message"),
                "created_at": video.get("created_at"),
                "completed_at": video.get("completed_at")
            }
        
        # Si la vidéo est en cours de traitement, vérifier son statut via l'API
        if video.get("task_id"):
            invideo_client = InVideoClient()
            
            # Vérifier le statut auprès d'InVideo
            status_result = invideo_client.get_video_status(video.get("task_id"))
            
            # Si le statut a changé, mettre à jour dans la base de données
            if status_result.get("status") != video.get("status"):
                update_data = {
                    "status": status_result.get("status")
                }
                
                # Ajouter d'autres données en fonction du statut
                if status_result.get("status") == "completed":
                    update_data.update({
                        "video_url": status_result.get("video_url"),
                        "thumbnail_url": status_result.get("thumbnail_url"),
                        "completed_at": datetime.utcnow()
                    })
                elif status_result.get("status") == "failed":
                    update_data.update({
                        "error_message": status_result.get("error_message"),
                        "completed_at": datetime.utcnow()
                    })
                
                # Mettre à jour la vidéo
                db.videos.update_one(
                    {"_id": ObjectId(video_id)},
                    {"$set": update_data}
                )
                
                # Si la vidéo est terminée, mettre également à jour le statut du script associé
                if status_result.get("status") in ["completed", "failed"]:
                    db.batch_scripts.update_one(
                        {"_id": ObjectId(video.get("script_id"))},
                        {"$set": {
                            "status": "video_generated" if status_result.get("status") == "completed" else "video_failed"
                        }}
                    )
                
                # Mettre à jour l'objet vidéo avec les nouvelles données
                video.update(update_data)
            
            # Retourner les données actualisées
            return {
                "id": str(video["_id"]),
                "title": video.get("title"),
                "status": video.get("status"),
                "progress": status_result.get("progress"),
                "video_url": video.get("video_url"),
                "thumbnail_url": video.get("thumbnail_url"),
                "error_message": video.get("error_message"),
                "created_at": video.get("created_at"),
                "completed_at": video.get("completed_at")
            }
        
        # Si pas de task_id, renvoyer simplement le statut actuel
        return {
            "id": str(video["_id"]),
            "title": video.get("title"),
            "status": video.get("status"),
            "created_at": video.get("created_at")
        }
        
    except Exception as e:
        print(f"Erreur lors de la vérification du statut de la vidéo {video_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Erreur lors de la vérification du statut: {str(e)}")

# Endpoint pour vérifier le statut global d'un lot et mettre à jour si nécessaire
@app.get("/videos/batch/{batch_id}/status")
async def check_batch_status(
    batch_id: str,
    current_user: UserInDB = Depends(get_current_user)
):
    db = get_db()
    
    try:
        # Vérifier que le lot existe et appartient à l'utilisateur
        batch = db.batch_videos.find_one({
            "_id": ObjectId(batch_id),
            "created_by": current_user.username
        })
        
        if not batch:
            raise HTTPException(status_code=404, detail="Lot non trouvé ou vous n'avez pas la permission de le consulter")
        
        # Si le lot n'est pas en cours de traitement, renvoyer simplement son statut
        if batch.get("status") != "processing":
            return {
                "id": batch_id,
                "status": batch.get("status"),
                "total_videos": batch.get("total_videos"),
                "processed_videos": batch.get("processed_videos"),
                "completed_videos": batch.get("completed_videos", 0),
                "failed_videos": batch.get("failed_videos", 0),
                "created_at": batch.get("created_at"),
                "start_processing_time": batch.get("start_processing_time"),
                "end_time": batch.get("end_time")
            }
        
        # Si le lot est en cours de traitement, vérifier le statut de toutes les vidéos
        videos = list(db.videos.find({"batch_id": batch_id}))
        
        total = len(videos)
        completed = 0
        failed = 0
        processing = 0
        
        # Compter les statuts
        for video in videos:
            if video.get("status") == "completed":
                completed += 1
            elif video.get("status") == "failed":
                failed += 1
            else:
                processing += 1
        
        # Déterminer le statut global du lot
        batch_status = "processing"
        
        # Si toutes les vidéos sont terminées (complétées ou échouées)
        if processing == 0:
            if failed == 0:
                batch_status = "completed"
            elif completed == 0:
                batch_status = "failed"
            else:
                batch_status = "partially_completed"
            
            # Mettre à jour le statut final dans la base de données
            db.batch_videos.update_one(
                {"_id": ObjectId(batch_id)},
                {"$set": {
                    "status": batch_status,
                    "processed_videos": total,
                    "completed_videos": completed,
                    "failed_videos": failed,
                    "end_time": datetime.utcnow()
                }}
            )
        else:
            # Mettre à jour les compteurs si le traitement est toujours en cours
            db.batch_videos.update_one(
                {"_id": ObjectId(batch_id)},
                {"$set": {
                    "processed_videos": total,
                    "completed_videos": completed,
                    "failed_videos": failed,
                    "last_updated_at": datetime.utcnow()
                }}
            )
        
        return {
            "id": batch_id,
            "status": batch_status,
            "total_videos": total,
            "processed_videos": total,
            "completed_videos": completed,
            "failed_videos": failed,
            "processing_videos": processing,
            "created_at": batch.get("created_at"),
            "start_processing_time": batch.get("start_processing_time"),
            "end_time": datetime.utcnow() if batch_status != "processing" else None
        }
        
    except Exception as e:
        print(f"Erreur lors de la vérification du statut du lot {batch_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Erreur lors de la vérification du statut: {str(e)}")

# Run the server
if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8001, reload=True)

def generate_article_content(prompt_type: str, theme: str, additional_content: Optional[str] = None, user: UserInDB = Depends(get_current_user)):
    # Use the user's OpenAI API key if available, otherwise fall back to the default
    openai_api_key = user.openai_api_key or os.getenv("OPENAI_API_KEY")
    openai.api_key = openai_api_key
    print(f"OPENAI_API_KEY: {openai_api_key}")
    # Get base prompt based on type
    base_prompts = {
        "cybersecurity": """Écris un article détaillé en français pour LinkedIn sous une forme rédactionnelle de type journalistique.

Instructions spécifiques:
- Ne pas numéroter les paragraphes
- Intégrer les titres de paragraphe naturellement dans le texte
- Mettre en gras les points importants
- Générer une accroche avec une question ouverte
- Inclure une conclusion impactante
- Suggérer une image professionnelle pertinente (description détaillée)
- Générer 5-7 hashtags pertinents en français
- Longueur: 800-1200 mots
- Style: professionnel mais accessible
- Ton: expert et pédagogue

Structure de la réponse:
===Titre===
[Titre de l'article]

===IMAGE SUGGÉRÉE===
[Description détaillée de l'image recommandée]

===ARTICLE===
[Contenu principal de l'article]

===HASHTAGS===
[Liste des hashtags]""",
        
        "digital": """Écris un article détaillé en français pour LinkedIn sous une forme rédactionnelle de type journalistique.

Instructions spécifiques:
- Ne pas numéroter les paragraphes
- Intégrer les titres de paragraphe naturellement dans le texte
- Mettre en gras les points importants
- Générer une accroche avec une question ouverte
- Inclure une conclusion impactante
- Suggérer une image professionnelle pertinente (description détaillée)
- Générer 5-7 hashtags pertinents en français
- Longueur: 800-1200 mots
- Style: professionnel mais accessible
- Ton: expert et pédagogue

Structure de la réponse:
===Titre===
[Titre de l'article]

===IMAGE SUGGÉRÉE===
[Description détaillée de l'image recommandée]

===ARTICLE===
[Contenu principal de l'article]

===HASHTAGS===
[Liste des hashtags]"""
    }
    
    # Get the prompt template or use default (cybersecurity)
    prompt_template = base_prompts.get(prompt_type, base_prompts["cybersecurity"])
    
    # Construct full prompt
    full_prompt = f"{prompt_template}\n\nThème spécifique à traiter: {theme}"
    if additional_content:
        full_prompt += f"\nContexte additionnel: {additional_content}"
    
    # Print the prompt being sent to OpenAI
    print("\n=== Prompt envoyé à ChatGPT ===")
    print(full_prompt)
    print("==============================\n")
    
    # Call OpenAI API
    response = openai.ChatCompletion.create(
        model="gpt-4o",
        messages=[
            {"role": "user", "content": full_prompt}
        ],
        temperature=0.7,
        max_tokens=2000
    )
    
    # Get the generated content
    generated_content = response.choices[0].message['content']
    
    # Print ChatGPT's response
    print("\n=== Réponse de ChatGPT ===")
    print(generated_content)
    print("===========================\n")
    
    return generated_content

def save_as_word_doc(title: str, content: str, output_path: str):
    doc = docx.Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(content)
    doc.save(output_path)
    return output_path

async def generate_image_from_prompt(prompt: str, openai_api_key: str):
    """
    Generate an image using OpenAI's DALL-E based on the provided prompt.
    Returns a tuple of (image_path, error_message)
    """
    try:
        openai.api_key = openai_api_key
        
        # Log the image generation attempt
        print(f"\n=== Generating image with DALL-E ===")
        print(f"Original prompt: {prompt[:100]}...")
        
        # Define possible styles with their corresponding prompts
        styles = {
            "flat_design": "Create a professional flat design illustration with clean lines and minimal details. Use bright, modern colors with simple geometric shapes. Style similar to modern corporate illustrations or infographics. No text.",
            
            "isometric": "Create a professional isometric illustration with clean 3D perspective. Use a cohesive color palette with subtle gradients and shadows for depth. Style similar to modern tech company illustrations. No text.",
            
            "conceptual": "Create a conceptual illustration with metaphorical elements. Use subtle symbolism and clean professional aesthetics. Style similar to high-quality business magazine illustrations. No text.",
            
            "digital_painting": "Create a polished digital painting with professional aesthetics. Use soft lighting, balanced composition, and a cohesive color scheme. Suitable for business or professional context. No text."
        }
        
        # Select a style based on the content of the prompt
        selected_style = "flat_design"  # Default style
        
        if any(tech_term in prompt.lower() for tech_term in ["digital", "technologie", "informatique", "cyber", "data", "internet", "réseau", "virtuel"]):
            selected_style = "isometric"
        elif any(concept_term in prompt.lower() for concept_term in ["concept", "idée", "stratégie", "croissance", "innovation", "leadership", "partenariat"]):
            selected_style = "conceptual"
        elif any(creative_term in prompt.lower() for creative_term in ["créatif", "artistique", "design", "moderne", "tendance"]):
            selected_style = "digital_painting"
            
        # Craft an enhanced prompt with the selected style
        style_prompt = styles[selected_style]
        enhanced_prompt = f"{style_prompt} Subject: {prompt}. High quality, professional image suitable for LinkedIn or business publications. Use a balanced composition with clear visual hierarchy."
        
        print(f"Selected style: {selected_style}")
        print(f"Enhanced prompt: {enhanced_prompt[:100]}...")
        
        # Generate image with DALL-E, using a higher quality model if available
        response = openai.Image.create(
            prompt=enhanced_prompt,
            n=1,  # Number of images to generate
            size="1024x1024",  # Image size
            response_format="url",  # Get URL instead of base64 to save bandwidth
            quality="hd",  # Request higher quality if available
            model="dall-e-3"  # Use the most recent DALL-E 3 model
        )
        
        image_url = response['data'][0]['url']
        
        # Download the image
        image_response = requests.get(image_url)
        if image_response.status_code != 200:
            return None, "Failed to download the generated image"
        
        # Create directory if it doesn't exist
        os.makedirs("generated_images", exist_ok=True)
        
        # Save the image with a unique filename
        image_filename = f"article_image_{uuid.uuid4()}.png"
        image_path = os.path.join("generated_images", image_filename)
        
        # Open and save the image using PIL
        image = Image.open(io.BytesIO(image_response.content))
        image.save(image_path)
        
        print(f"✅ Image successfully generated with style '{selected_style}' and saved to {image_path}")
        
        return image_path, None, selected_style
        
    except Exception as e:
        error_message = f"Error generating image: {str(e)}"
        print(f"❌ {error_message}")
        return None, error_message, None

def send_email(to_email: str, subject: str, content: str, attachment_paths: Optional[List[str]] = None, is_html: bool = False):
    """
    Send an email with optional attachments.
    
    Args:
        to_email: Recipient email address
        subject: Email subject
        content: Email content (plain text or HTML)
        attachment_paths: List of file paths to attach
        is_html: Whether the content is HTML (default: False)
    
    Returns:
        bool: True if email was sent successfully, False otherwise
    """
    try:
        # Email configuration
        smtp_server = "smtp.gmail.com"
        smtp_port = 587
        sender_email = os.getenv("EMAIL_ADDRESS")
        sender_password = os.getenv("EMAIL_PASSWORD")
        
        if not sender_email or not sender_password:
            logger.error("EMAIL_ADDRESS or EMAIL_PASSWORD environment variables not set")
            return False
        
        # Create message
        msg = MIMEMultipart()
        msg['From'] = sender_email
        msg['To'] = to_email
        msg['Subject'] = subject
        
        # Add body
        content_type = 'html' if is_html else 'plain'
        msg.attach(MIMEText(content, content_type))
        
        # Add attachments if provided
        if attachment_paths:
            for attachment_path in attachment_paths:
                if attachment_path and os.path.exists(attachment_path):
                    try:
                        with open(attachment_path, 'rb') as f:
                            part = MIMEApplication(f.read(), _subtype=attachment_path.split('.')[-1])
                            part.add_header('Content-Disposition', 'attachment', filename=os.path.basename(attachment_path))
                            msg.attach(part)
                    except Exception as attach_err:
                        logger.error(f"Error attaching file {attachment_path}: {str(attach_err)}")
        
        # Send email
        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls()
            server.login(sender_email, sender_password)
            server.send_message(msg)
            logger.info(f"Email sent successfully to {to_email}")
            return True
    except Exception as e:
        logger.error(f"Error sending email to {to_email}: {str(e)}")
        return False

@app.get("/images/{image_name}")
async def get_image(image_name: str):
    image_path = os.path.join("generated_images", image_name)
    if os.path.exists(image_path):
        return FileResponse(image_path)
    raise HTTPException(status_code=404, detail="Image not found")

# Fonction pour initialiser les prompts vidéo par défaut
async def initialize_default_video_prompts():
    db = get_db()
    
    # Vérifier si des prompts système existent déjà
    system_prompts_count = db.video_prompts.count_documents({"is_system": True})
    
    if system_prompts_count > 0:
        print(f"✅ {system_prompts_count} prompts vidéo système déjà initialisés")
        return
    
    # Définir les prompts système par défaut
    default_prompts = [
        {
            "name": "Promotionnel - Produit",
            "description": "Prompt pour des vidéos promotionnelles de produits",
            "prompt_template": "Crée un script de vidéo promotionnelle pour [THEME]. Le script doit présenter les bénéfices clés, inclure un appel à l'action clair, et durer environ 60 secondes.",
            "category": "commercial",
            "tags": ["promotionnel", "produit", "vente"],
            "is_system": True,
            "created_by": "system",
            "created_at": datetime.utcnow()
        },
        {
            "name": "Explicatif - Tutoriel",
            "description": "Prompt pour des vidéos explicatives et tutoriels",
            "prompt_template": "Rédige un script de tutoriel sur [THEME]. Le script doit être structuré étape par étape, expliquer clairement chaque concept et se terminer par un résumé des points clés.",
            "category": "éducation",
            "tags": ["tutoriel", "explication", "howto"],
            "is_system": True,
            "created_by": "system",
            "created_at": datetime.utcnow()
        },
        {
            "name": "Actualité - Tech",
            "description": "Prompt pour des vidéos d'actualité technologique",
            "prompt_template": "Écris un script de vidéo d'actualité sur [THEME]. Le script doit présenter les dernières nouvelles, fournir un contexte et une analyse, et se terminer par les implications futures.",
            "category": "actualité",
            "tags": ["tech", "news", "analyse"],
            "is_system": True,
            "created_by": "system",
            "created_at": datetime.utcnow()
        }
    ]
    
    # Insérer les prompts par défaut
    try:
        db.video_prompts.insert_many(default_prompts)
        print(f"✅ {len(default_prompts)} prompts vidéo système initialisés avec succès")
    except Exception as e:
        print(f"❌ Erreur lors de l'initialisation des prompts vidéo système: {str(e)}")

@app.get("/admin/stats")
async def get_admin_stats(current_user: UserInDB = Depends(get_current_admin_user)):
    db = get_db()
    
    # Get total counts
    total_users = await db.users.count_documents({})
    total_articles = await db.articles.count_documents({})
    total_videos = await db.videos.count_documents({})
    
    # Get active users (users who have created content in the last 30 days)
    thirty_days_ago = datetime.utcnow() - timedelta(days=30)
    active_users = len(await db.users.distinct('email', {
        '$or': [
            {'created_at': {'$gte': thirty_days_ago}},
            {'last_login': {'$gte': thirty_days_ago}}
        ]
    }))
    
    return {
        "totalUsers": total_users,
        "totalArticles": total_articles,
        "totalVideos": total_videos,
        "activeUsers": active_users
    }

@app.get("/admin/articles")
async def get_all_articles(current_user: UserInDB = Depends(get_current_admin_user)):
    """Get all articles for admin view."""
    db = get_db()
    cursor = db.articles.find({})
    articles = list(cursor)
    return [
        {
            "id": str(article["_id"]),
            "uid": article.get("uid", ""),
            "title": article.get("title", ""),
            "content": article.get("content", {}),
            "author": article.get("author", ""),
            "created_at": article.get("created_at", datetime.now())
        }
        for article in articles
    ]

@app.delete("/admin/articles/{article_id}")
async def delete_article(article_id: str, current_user: UserInDB = Depends(get_current_admin_user)):
    db = get_db()
    
    # Check if article exists
    article = await db.articles.find_one({"_id": ObjectId(article_id)})
    if not article:
        raise HTTPException(status_code=404, detail="Article not found")
    
    # Delete article
    await db.articles.delete_one({"_id": ObjectId(article_id)})
    
    return {"status": "success"}

@app.get("/admin/videos")
async def get_all_videos(current_user: UserInDB = Depends(get_current_admin_user)):
    """Get all videos for admin view."""
    db = get_db()
    cursor = db.videos.find({})
    videos = list(cursor)
    return [
        {
            "id": str(video["_id"]),
            "title": video.get("title", ""),
            "status": video.get("status", "unknown"),
            "video_url": video.get("video_url", None),
            "thumbnail_url": video.get("thumbnail_url", None),
            "author": video.get("author", ""),
            "created_at": video.get("created_at", datetime.now())
        }
        for video in videos
    ]

@app.delete("/admin/videos/{video_id}")
async def delete_video(video_id: str, current_user: UserInDB = Depends(get_current_admin_user)):
    db = get_db()
    
    # Check if video exists
    video = await db.videos.find_one({"_id": ObjectId(video_id)})
    if not video:
        raise HTTPException(status_code=404, detail="Video not found")
    
    # Delete video
    await db.videos.delete_one({"_id": ObjectId(video_id)})
    
    return {"status": "success"} 

@app.get("/admin/users")
async def get_all_users(current_user: UserInDB = Depends(get_current_admin_user)):
    """Get all users for admin view."""
    db = get_db()
    cursor = db.users.find({})
    users = list(cursor)
    return [
        {
            "id": str(user["_id"]),
            "username": user.get("username", ""),
            "email": user.get("email", ""),
            "role": user.get("role", "user"),
            "created_at": user.get("created_at", datetime.now())
        }
        for user in users
    ]

@app.put("/admin/users/{user_id}/role")
async def update_user_role(
    user_id: str, 
    role: str, 
    current_user: UserInDB = Depends(get_current_admin_user)
):
    """Update a user's role. Only accessible by admins."""
    db = get_db()
    
    # Verify the role is valid
    if role not in ["user", "client", "admin"]:
        raise HTTPException(status_code=400, detail="Invalid role")
    
    # Don't allow modifying owner accounts
    target_user = db.users.find_one({"_id": ObjectId(user_id)})
    if not target_user:
        raise HTTPException(status_code=404, detail="User not found")
    
    if target_user.get("role") == "owner":
        raise HTTPException(status_code=403, detail="Cannot modify owner account")
    
    # Update the user's role
    result = db.users.update_one(
        {"_id": ObjectId(user_id)},
        {"$set": {"role": role}}
    )
    
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="User not found or no changes made")
    
    return {"status": "success", "message": f"User role updated to {role}"}

# Add new model for bulk article generation
class ThemeData(BaseModel):
    subject: str
    theme_name: str
    custom_prompt: Optional[str] = None

class BulkArticleGenerate(BaseModel):
    themes_data: List[ThemeData] = []
    themes: List[str] = []  # Keep for backward compatibility
    email_to: EmailStr
    prompt_type: str = "cybersecurity"

@app.post("/articles/bulk", status_code=status.HTTP_202_ACCEPTED)
async def create_bulk_articles(bulk_request: BulkArticleGenerate, current_user: UserInDB = Depends(get_current_user)):
    """
    Generate multiple articles based on themes provided in a CSV file.
    """
    try:
        # Create a record in the database to track this bulk generation
        bulk_id = str(uuid.uuid4())
        now = datetime.now()
        
        # Extract user ID correctly from UserInDB object
        user_id = str(current_user.id) if hasattr(current_user, 'id') else str(current_user._id if hasattr(current_user, '_id') else current_user['_id'])
        
        # Determine the total number of articles to process
        total_articles = len(bulk_request.themes_data) if bulk_request.themes_data else len(bulk_request.themes)
        
        bulk_record = {
            "id": bulk_id,
            "user_id": user_id,
            "email": bulk_request.email_to,
            "prompt_type": bulk_request.prompt_type,
            "total_articles": total_articles,
            "processed_articles": 0,
            "status": "processing",
            "created_at": now,
            "articles": []
        }
        
        db = get_db()
        # Use insert_one without await for synchronous operations
        db.article_batches.insert_one(bulk_record)
        
        # Start a background task to generate articles
        # If we have themes_data with custom prompts, use that
        if bulk_request.themes_data:
            asyncio.create_task(process_bulk_articles_with_custom_prompts(
                bulk_id=bulk_id,
                themes_data=bulk_request.themes_data,
                email_to=bulk_request.email_to,
                default_prompt_type=bulk_request.prompt_type,
                user=current_user
            ))
        else:
            # Backward compatibility for old format
            asyncio.create_task(process_bulk_articles(
                bulk_id=bulk_id,
                themes=bulk_request.themes,
                email_to=bulk_request.email_to,
                prompt_type=bulk_request.prompt_type,
                user=current_user
            ))
        
        return {"message": "Bulk article generation started", "bulk_id": bulk_id, "processed": total_articles}
    except Exception as e:
        logger.error(f"Error starting bulk article generation: {e}")
        raise HTTPException(status_code=500, detail=f"An error occurred: {str(e)}")

async def process_bulk_articles_with_custom_prompts(bulk_id: str, themes_data: List[ThemeData], email_to: str, default_prompt_type: str, user: UserInDB):
    try:
        total = len(themes_data)
        processed = 0
        articles = []
        db = get_db()
        
        # Extract username correctly from UserInDB object
        username = user.username if hasattr(user, 'username') else user['username']
        
        # Retrieve the user's OpenAI API key
        user_record = db.users.find_one({"email": user.email})
        openai_api_key = user_record.get("openai_api_key") or os.getenv("OPENAI_API_KEY")
        print(f"OpenAI API key: {openai_api_key}")
        
        for theme_data in themes_data:
            try:
                subject = theme_data.subject
                theme_name = theme_data.theme_name
                custom_prompt = theme_data.custom_prompt
                
                # Set OpenAI API key
                openai.api_key = openai_api_key
                
                # Generate article content - use custom prompt if available
                if custom_prompt:
                    # Call with custom prompt instead of predefined prompt type
                    generated_content = generate_article_content_with_custom_prompt(
                        theme=subject,
                        custom_prompt=custom_prompt,
                        openai_api_key=openai_api_key
                    )
                else:
                    # Fallback to default prompt type
                    generated_content = generate_article_content(
                        prompt_type=default_prompt_type,
                        theme=subject
                    )
                
                # Parse the generated content (same as in process_bulk_articles)
                title = ""
                image_suggestion = ""
                article_content = ""
                hashtags = ""
                
                # Improved parsing logic
                if "===Titre===" in generated_content:
                    title_parts = generated_content.split("===Titre===")
                    if len(title_parts) > 1:
                        next_section = "===IMAGE SUGGÉRÉE===" if "===IMAGE SUGGÉRÉE===" in title_parts[1] else (
                                       "===ARTICLE===" if "===ARTICLE===" in title_parts[1] else 
                                       "===HASHTAGS===" if "===HASHTAGS===" in title_parts[1] else "")
                        
                        if next_section:
                            title = title_parts[1].split(next_section)[0].strip()
                        else:
                            title = title_parts[1].strip()
                
                if "===IMAGE SUGGÉRÉE===" in generated_content:
                    img_parts = generated_content.split("===IMAGE SUGGÉRÉE===")
                    if len(img_parts) > 1:
                        next_section = "===ARTICLE===" if "===ARTICLE===" in img_parts[1] else (
                                       "===HASHTAGS===" if "===HASHTAGS===" in img_parts[1] else "")
                        
                        if next_section:
                            image_suggestion = img_parts[1].split(next_section)[0].strip()
                        else:
                            image_suggestion = img_parts[1].strip()
                
                if "===ARTICLE===" in generated_content:
                    article_parts = generated_content.split("===ARTICLE===")
                    if len(article_parts) > 1:
                        next_section = "===HASHTAGS===" if "===HASHTAGS===" in article_parts[1] else ""
                        
                        if next_section:
                            article_content = article_parts[1].split(next_section)[0].strip()
                        else:
                            article_content = article_parts[1].strip()
                
                if "===HASHTAGS===" in generated_content:
                    hashtag_parts = generated_content.split("===HASHTAGS===")
                    if len(hashtag_parts) > 1:
                        hashtags = hashtag_parts[1].strip()
                
                # If no title was found in the content, use the theme
                if not title:
                    title = f"Article sur {subject}"
                
                # Create structured content dictionary
                structured_content = {
                    "title": title or f"Article sur {subject}",
                    "image_suggestion": image_suggestion or "Aucune suggestion d'image fournie",
                    "article_content": article_content or "Contenu de l'article non disponible",
                    "hashtags": hashtags or "Aucun hashtag fourni",
                    "raw_content": generated_content,
                }
                
                # Generate a unique identifier for the article
                uid = str(uuid.uuid4())
                
                # Create article object
                article = {
                    "uid": uid,
                    "title": structured_content["title"],
                    "category": default_prompt_type if not custom_prompt else "custom",
                    "summary": subject,
                    "content": structured_content,
                    "author": username,
                    "user_email": user.email if hasattr(user, 'email') else user.get('email', ''),
                    "created_at": datetime.now(),
                    "theme_name": theme_name,  # Store the theme name
                    "used_custom_prompt": bool(custom_prompt)  # Track if custom prompt was used
                }
                
                # Save to database - use without await for synchronous operations
                result = db.articles.insert_one(article)
                article["id"] = str(result.inserted_id)
                articles.append(article)
                
                # Try to generate an image if an image suggestion is provided
                if structured_content["image_suggestion"] and structured_content["image_suggestion"] != "Aucune suggestion d'image fournie":
                    try:
                        # Generate image
                        image_path, error, image_style = await generate_image_from_prompt(structured_content["image_suggestion"], openai_api_key)
                        
                        # Update article with image information
                        if image_path:
                            structured_content["image_path"] = image_path
                            structured_content["image_style"] = image_style
                            
                            # Update the article in the database - use without await
                            db.articles.update_one(
                                {"_id": result.inserted_id},
                                {"$set": {"content": structured_content}}
                            )
                    except Exception as img_error:
                        logger.error(f"Error generating image for article {uid}: {img_error}")
                
                # Generate Word document
                if structured_content["title"] and structured_content["article_content"]:
                    try:
                        # Create a filename based on the title
                        safe_title = ''.join(c if c.isalnum() or c in ' -_' else '_' for c in structured_content["title"])
                        filename = f"{default_prompt_type}_{safe_title[:30]}.docx"
                        output_path = f"temp/{filename}"
                        
                        # Ensure the temp directory exists
                        os.makedirs("temp", exist_ok=True)
                        
                        # Create Word document with article content
                        full_content = f"Title: {structured_content['title']}\n\n"
                        if structured_content["image_suggestion"] != "Aucune suggestion d'image fournie":
                            full_content += f"Image Suggestion: {structured_content['image_suggestion']}\n\n"
                        full_content += structured_content["article_content"]
                        if structured_content["hashtags"] != "Aucun hashtag fourni":
                            full_content += f"\n\nHashtags: {structured_content['hashtags']}"
                        
                        save_as_word_doc(structured_content["title"], full_content, output_path)
                        
                        # Prepare attachment list for email
                        attachments = [output_path]
                        
                        # Send email with article and attachment
                        email_subject = f"Article généré: {structured_content['title']}"
                        email_content = f"""
                        <h1>Votre article a été généré avec succès</h1>
                        <p><strong>Titre:</strong> {structured_content['title']}</p>
                        <p><strong>Thème:</strong> {subject}</p>
                        <p>Vous trouverez l'article complet en pièce jointe.</p>
                        <hr>
                        <h2>Aperçu de l'article:</h2>
                        <p>{structured_content['article_content'][:300]}...</p>
                        <hr>
                        <p><em>Service de génération automatique d'articles</em></p>
                        """
                        send_email(
                            to_email=email_to,
                            subject=email_subject,
                            content=email_content,
                            attachment_paths=attachments,
                            is_html=True
                        )
                        
                        # Delete the temporary file after sending
                        try:
                            os.remove(output_path)
                        except:
                            pass
                    except Exception as doc_error:
                        logger.error(f"Error creating or sending document for article {uid}: {doc_error}")
                
                # Update processed count
                processed += 1
                
                # Update batch record with progress - use without await
                db.article_batches.update_one(
                    {"id": bulk_id},
                    {
                        "$set": {
                            "processed_articles": processed,
                            "status": "processing" if processed < total else "completed"
                        },
                        "$push": {"articles": article["id"]}
                    }
                )
                
            except Exception as article_error:
                logger.error(f"Error processing article with theme '{subject}': {article_error}")
        
        # Send a completion email
        if processed > 0:
            completion_subject = f"Génération d'articles terminée - {processed}/{total} articles"
            completion_content = f"""
            <h1>Génération d'articles terminée</h1>
            <p>Nous avons généré <strong>{processed}</strong> articles sur <strong>{total}</strong> demandés.</p>
            <p>Vous avez reçu chaque article dans un email séparé avec une pièce jointe au format Word.</p>
            <hr>
            <p><em>Service de génération automatique d'articles</em></p>
            """
            send_email(
                to_email=email_to,
                subject=completion_subject,
                content=completion_content,
                is_html=True
            )
        
        # Mark batch as completed - use without await
        db.article_batches.update_one(
            {"id": bulk_id},
            {
                "$set": {
                    "status": "completed",
                    "completed_at": datetime.now()
                }
            }
        )
        
    except Exception as e:
        logger.error(f"Error processing bulk articles: {e}")
        # Mark batch as failed - use without await
        db.get_db().article_batches.update_one(
            {"id": bulk_id},
            {
                "$set": {
                    "status": "failed",
                    "completed_at": datetime.now(),
                    "error": str(e)
                }
            }
        )

# Function to generate article content with a custom prompt
def generate_article_content_with_custom_prompt(theme: str, custom_prompt: str, openai_api_key: str):
    # Set OpenAI API key
    openai.api_key = openai_api_key
    
    # Ensure the custom prompt includes the required output struct
    structure_appendix = """
Structure de la réponse (IMPORTANT, respectez ce format exactement):
===Titre===
[Titre de l'article]

===IMAGE SUGGÉRÉE===
[Description détaillée de l'image recommandée]

===ARTICLE===
[Contenu principal de l'article]

===HASHTAGS===
[Liste des hashtags]
"""
    
    # Add the structure if it's not already included
    if "===Titre===" not in custom_prompt:
        custom_prompt += structure_appendix
    
    # Construct full prompt
    full_prompt = f"{custom_prompt}\n\nThème spécifique à traiter: {theme}"
    
    # Print the prompt being sent to OpenAI
    print("\n=== Prompt personnalisé envoyé à ChatGPT ===")
    print(full_prompt)
    print("==========================================\n")
    
    # Call OpenAI API
    response = openai.ChatCompletion.create(
        model="gpt-4o",
        messages=[
            {"role": "user", "content": full_prompt}
        ],
        temperature=0.7,
        max_tokens=2000
    )
    
    # Get the generated content
    generated_content = response.choices[0].message['content']
    
    # Print ChatGPT's response
    print("\n=== Réponse de ChatGPT (prompt personnalisé) ===")
    print(generated_content)
    print("==============================================\n")
    
    return generated_content

# Status model for bulk article generation
class BulkArticleStatus(BaseModel):
    id: str
    user_id: str
    email: str
    prompt_type: str
    total_articles: int
    processed_articles: int
    status: str
    created_at: datetime
    completed_at: Optional[datetime] = None
    error: Optional[str] = None

@app.get("/articles/bulk/{bulk_id}/status", response_model=BulkArticleStatus)
async def get_bulk_article_status(bulk_id: str, current_user: UserInDB = Depends(get_current_user)):
    """
    Get the status of a bulk article generation process
    """
    try:
        # Get the batch record from the database
        db = get_db()
        # Use find_one without await for synchronous operation
        batch = db.article_batches.find_one({"id": bulk_id})
        
        if not batch:
            raise HTTPException(status_code=404, detail="Bulk generation not found")
        
        # Get user ID safely from UserInDB object
        user_id = str(current_user.id) if hasattr(current_user, 'id') else str(current_user._id if hasattr(current_user, '_id') else current_user['_id'])
        
        # Get user role safely
        user_role = current_user.role if hasattr(current_user, 'role') else current_user.get('role', 'user')
        
        # Check if the user is authorized to view this batch
        if str(batch["user_id"]) != user_id and user_role not in ["admin", "owner"]:
            raise HTTPException(status_code=403, detail="Not authorized to view this batch")
        
        # Return the batch status
        return batch
    except Exception as e:
        logger.error(f"Error getting bulk article status: {e}")
        raise HTTPException(status_code=500, detail=f"An error occurred: {str(e)}")

@app.get("/articles/bulks", response_model=List[BulkArticleStatus])
async def get_bulk_articles(current_user: UserInDB = Depends(get_current_user)):
    """
    Get all bulk article generations for the current user
    """
    try:
        # Get user ID safely from UserInDB object
        user_id = str(current_user.id) if hasattr(current_user, 'id') else str(current_user._id if hasattr(current_user, '_id') else current_user['_id'])
        
        # Get the batch records from the database
        db = get_db()
        # Use find without await for synchronous operation
        cursor = db.article_batches.find({"user_id": user_id})
        batches = list(cursor)
        
        # Sort by created_at descending (newest first)
        batches.sort(key=lambda x: x.get("created_at", datetime.min), reverse=True)
        
        return batches
    except Exception as e:
        logger.error(f"Error getting bulk articles: {e}")
        raise HTTPException(status_code=500, detail=f"An error occurred: {str(e)}")

@app.get("/articles/bulk/{bulk_id}/articles", response_model=List[ArticleResponse])
async def get_bulk_articles_content(bulk_id: str, current_user: UserInDB = Depends(get_current_user)):
    """
    Get all articles in a specific bulk generation
    """
    try:
        db = get_db()
        
        # Get the batch record
        batch = db.article_batches.find_one({"id": bulk_id})
        if not batch:
            raise HTTPException(status_code=404, detail="Bulk generation not found")
        
        # Get user ID safely from UserInDB object
        user_id = str(current_user.id) if hasattr(current_user, 'id') else str(current_user._id if hasattr(current_user, '_id') else current_user['_id'])
        
        # Get user role safely
        user_role = current_user.role if hasattr(current_user, 'role') else current_user.get('role', 'user')
        
        # Check if the user is authorized to view this batch
        if str(batch["user_id"]) != user_id and user_role not in ["admin", "owner"]:
            raise HTTPException(status_code=403, detail="Not authorized to view this batch")
        
        # Get all articles in this batch
        articles = []
        article_ids = batch.get("articles", [])
        
        # Convert string IDs to ObjectId
        from bson import ObjectId
        object_ids = [ObjectId(id) for id in article_ids]
        
        # Fetch articles by their IDs
        if object_ids:
            for article in db.articles.find({"_id": {"$in": object_ids}}):
                article["id"] = str(article.pop("_id"))
                articles.append(article)
        
        return articles
    except Exception as e:
        logger.error(f"Error getting bulk articles content: {e}")
        raise HTTPException(status_code=500, detail=f"An error occurred: {str(e)}")

@app.delete("/articles/{article_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_user_article(article_id: str, current_user: UserInDB = Depends(get_current_user)):
    """
    Delete an article by ID
    """
    try:
        db = get_db()
        
        # Convert string ID to ObjectId
        from bson import ObjectId
        object_id = ObjectId(article_id)
        
        # First check if the article exists and belongs to the current user
        article = db.articles.find_one({"_id": object_id})
        if not article:
            raise HTTPException(status_code=404, detail="Article not found")
        
        # Get user email and role safely
        user_email = current_user.email if hasattr(current_user, 'email') else current_user.get('email', '')
        user_role = current_user.role if hasattr(current_user, 'role') else current_user.get('role', 'user')
        
        # Check if the user is authorized to delete this article
        if article.get("user_email") != user_email and user_role not in ["admin", "owner"]:
            raise HTTPException(status_code=403, detail="Not authorized to delete this article")
        
        # Delete the article
        result = db.articles.delete_one({"_id": object_id})
        
        if result.deleted_count == 0:
            raise HTTPException(status_code=404, detail="Article not found")
        
        # Also remove the article ID from any bulk article batches
        db.article_batches.update_many(
            {"articles": article_id},
            {"$pull": {"articles": article_id}}
        )
        
        return Response(status_code=status.HTTP_204_NO_CONTENT)
    except Exception as e:
        logger.error(f"Error deleting article: {e}")
        raise HTTPException(status_code=500, detail=f"An error occurred: {str(e)}")

# User update models
class UserUpdate(BaseModel):
    username: str

class PasswordChange(BaseModel):
    current_password: str
    new_password: str
    logout_all_sessions: bool = False  # Optional flag to logout from all sessions

@app.put("/users/update", status_code=status.HTTP_200_OK)
async def update_user(user_update: UserUpdate, current_user: UserInDB = Depends(get_current_user)):
    """
    Update the current user's profile information
    """
    db = get_db()
    
    # Check if the username is already taken by another user
    existing_user = db.users.find_one({"username": user_update.username, "email": {"$ne": current_user.email}})
    if existing_user:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Username already taken"
        )
    
    # Update the user information
    result = db.users.update_one(
        {"email": current_user.email},
        {"$set": {"username": user_update.username}}
    )
    
    if result.modified_count == 0:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="User update failed"
        )
    
    return {"message": "User updated successfully"}

@app.post("/users/change-password", status_code=status.HTTP_200_OK)
async def change_password(password_change: PasswordChange, current_user: UserInDB = Depends(get_current_user)):
    """
    Change the current user's password
    """
    db = get_db()
    
    # Verify the current password
    if not verify_password(password_change.current_password, current_user.hashed_password):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Incorrect current password"
        )
    
    # Generate new hashed password
    new_hashed_password = get_password_hash(password_change.new_password)
    
    update_data = {"hashed_password": new_hashed_password}
    
    # If user requested to logout from all sessions, add a token blacklist timestamp
    if password_change.logout_all_sessions:
        update_data["password_changed_at"] = datetime.utcnow()
    
    # Update the password in the database
    result = db.users.update_one(
        {"email": current_user.email},
        {"$set": update_data}
    )
    
    if result.modified_count == 0:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Password update failed"
        )
    
    return {"message": "Password changed successfully", "logout_required": password_change.logout_all_sessions}

class UserTheme(BaseModel):
    name: str
    prompt: str

@app.post("/users/themes", status_code=status.HTTP_201_CREATED)
async def add_user_theme(theme: UserTheme, current_user: UserInDB = Depends(get_current_user)):
    """
    Add a new theme to the user's profile
    """
    db = get_db()
    
    # Create theme document
    theme_doc = {
        "name": theme.name,
        "prompt": theme.prompt,
        "user_id": str(current_user.id) if hasattr(current_user, 'id') else str(current_user._id if hasattr(current_user, '_id') else current_user['_id']),
        "created_at": datetime.utcnow()
    }
    
    # Insert the theme
    result = db.user_themes.insert_one(theme_doc)
    
    return {
        "id": str(result.inserted_id),
        "name": theme.name,
        "prompt": theme.prompt,
        "created_at": theme_doc["created_at"]
    }

@app.get("/users/themes", response_model=List[dict])
async def get_user_themes(current_user: UserInDB = Depends(get_current_user)):
    """
    Get all themes for the current user
    """
    db = get_db()
    
    # Get user ID
    user_id = str(current_user.id) if hasattr(current_user, 'id') else str(current_user._id if hasattr(current_user, '_id') else current_user['_id'])
    
    # Find themes
    themes = list(db.user_themes.find({"user_id": user_id}))
    
    # Format themes for response
    response = []
    for theme in themes:
        response.append({
            "id": str(theme["_id"]),
            "name": theme["name"],
            "prompt": theme["prompt"],
            "created_at": theme["created_at"]
        })
    
    return response

@app.put("/users/themes/{theme_id}", status_code=status.HTTP_200_OK)
async def update_user_theme(theme_id: str, theme: UserTheme, current_user: UserInDB = Depends(get_current_user)):
    """
    Update a user theme
    """
    db = get_db()
    
    # Get user ID
    user_id = str(current_user.id) if hasattr(current_user, 'id') else str(current_user._id if hasattr(current_user, '_id') else current_user['_id'])
    
    # Update the theme
    result = db.user_themes.update_one(
        {"_id": ObjectId(theme_id), "user_id": user_id},
        {"$set": {
            "name": theme.name,
            "prompt": theme.prompt,
            "updated_at": datetime.utcnow()
        }}
    )
    
    if result.modified_count == 0:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Theme not found or you don't have permission to modify it"
        )
    
    return {"message": "Theme updated successfully"}

@app.delete("/users/themes/{theme_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_user_theme(theme_id: str, current_user: UserInDB = Depends(get_current_user)):
    """
    Delete a user theme
    """
    db = get_db()
    
    # Get user ID
    user_id = str(current_user.id) if hasattr(current_user, 'id') else str(current_user._id if hasattr(current_user, '_id') else current_user['_id'])
    
    # Delete the theme
    result = db.user_themes.delete_one({"_id": ObjectId(theme_id), "user_id": user_id})
    
    if result.deleted_count == 0:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Theme not found or you don't have permission to delete it"
        )

# Add back the original process_bulk_articles function for backward compatibility

async def process_bulk_articles(bulk_id: str, themes: List[str], email_to: str, prompt_type: str, user: UserInDB):
    """
    Background task to process multiple article generation.
    """
    try:
        total = len(themes)
        processed = 0
        articles = []
        db = get_db()
        
        # Extract username correctly from UserInDB object
        username = user.username if hasattr(user, 'username') else user['username']
        
        for theme in themes:
            try:
                # Generate article content
                generated_content = generate_article_content(prompt_type, theme)
                
                # Parse the generated content
                title = ""
                image_suggestion = ""
                article_content = ""
                hashtags = ""
                
                # Improved parsing logic
                if "===Titre===" in generated_content:
                    title_parts = generated_content.split("===Titre===")
                    if len(title_parts) > 1:
                        next_section = "===IMAGE SUGGÉRÉE===" if "===IMAGE SUGGÉRÉE===" in title_parts[1] else (
                                       "===ARTICLE===" if "===ARTICLE===" in title_parts[1] else 
                                       "===HASHTAGS===" if "===HASHTAGS===" in title_parts[1] else "")
                        
                        if next_section:
                            title = title_parts[1].split(next_section)[0].strip()
                        else:
                            title = title_parts[1].strip()
                
                if "===IMAGE SUGGÉRÉE===" in generated_content:
                    img_parts = generated_content.split("===IMAGE SUGGÉRÉE===")
                    if len(img_parts) > 1:
                        next_section = "===ARTICLE===" if "===ARTICLE===" in img_parts[1] else (
                                       "===HASHTAGS===" if "===HASHTAGS===" in img_parts[1] else "")
                        
                        if next_section:
                            image_suggestion = img_parts[1].split(next_section)[0].strip()
                        else:
                            image_suggestion = img_parts[1].strip()
                
                if "===ARTICLE===" in generated_content:
                    article_parts = generated_content.split("===ARTICLE===")
                    if len(article_parts) > 1:
                        next_section = "===HASHTAGS===" if "===HASHTAGS===" in article_parts[1] else ""
                        
                        if next_section:
                            article_content = article_parts[1].split(next_section)[0].strip()
                        else:
                            article_content = article_parts[1].strip()
                
                if "===HASHTAGS===" in generated_content:
                    hashtag_parts = generated_content.split("===HASHTAGS===")
                    if len(hashtag_parts) > 1:
                        hashtags = hashtag_parts[1].strip()
                
                # If no title was found in the content, use the theme
                if not title:
                    title = f"Article sur {theme}"
                
                # Create structured content dictionary
                structured_content = {
                    "title": title or f"Article sur {theme}",
                    "image_suggestion": image_suggestion or "Aucune suggestion d'image fournie",
                    "article_content": article_content or "Contenu de l'article non disponible",
                    "hashtags": hashtags or "Aucun hashtag fourni",
                    "raw_content": generated_content,
                }
                
                # Generate a unique identifier for the article
                uid = str(uuid.uuid4())
                
                # Create article object
                article = {
                    "uid": uid,
                    "title": structured_content["title"],
                    "category": prompt_type,
                    "summary": theme,
                    "content": structured_content,
                    "author": username,
                    "user_email": user.email if hasattr(user, 'email') else user.get('email', ''),
                    "created_at": datetime.now(),
                }
                
                # Save to database - use without await for synchronous operations
                result = db.articles.insert_one(article)
                article["id"] = str(result.inserted_id)
                articles.append(article)
                
                # Set OpenAI API key
                user_record = db.users.find_one({"email": user.email})
                openai_api_key = user_record.get("openai_api_key") or os.getenv("OPENAI_API_KEY")
                print(f"OpenAI API key: {openai_api_key}")

                openai.api_key = openai_api_key
                
                # Try to generate an image if an image suggestion is provided
                if structured_content["image_suggestion"] and structured_content["image_suggestion"] != "Aucune suggestion d'image fournie":
                    try:
                        # Generate image
                        image_path, error, image_style = await generate_image_from_prompt(structured_content["image_suggestion"], openai_api_key)
                        
                        # Update article with image information
                        if image_path:
                            structured_content["image_path"] = image_path
                            structured_content["image_style"] = image_style
                            
                            # Update the article in the database - use without await
                            db.articles.update_one(
                                {"_id": result.inserted_id},
                                {"$set": {"content": structured_content}}
                            )
                    except Exception as img_error:
                        logger.error(f"Error generating image for article {uid}: {img_error}")
                
                # Generate Word document
                if structured_content["title"] and structured_content["article_content"]:
                    try:
                        # Create a filename based on the title
                        safe_title = ''.join(c if c.isalnum() or c in ' -_' else '_' for c in structured_content["title"])
                        filename = f"{prompt_type}_{safe_title[:30]}.docx"
                        output_path = f"temp/{filename}"
                        
                        # Ensure the temp directory exists
                        os.makedirs("temp", exist_ok=True)
                        
                        # Create Word document with article content
                        full_content = f"Title: {structured_content['title']}\n\n"
                        if structured_content["image_suggestion"] != "Aucune suggestion d'image fournie":
                            full_content += f"Image Suggestion: {structured_content['image_suggestion']}\n\n"
                        full_content += structured_content["article_content"]
                        if structured_content["hashtags"] != "Aucun hashtag fourni":
                            full_content += f"\n\nHashtags: {structured_content['hashtags']}"
                        
                        save_as_word_doc(structured_content["title"], full_content, output_path)
                        
                        # Prepare attachment list for email
                        attachments = [output_path]
                        
                        # Send email with article and attachment
                        email_subject = f"Article généré: {structured_content['title']}"
                        email_content = f"""
                        <h1>Votre article a été généré avec succès</h1>
                        <p><strong>Titre:</strong> {structured_content['title']}</p>
                        <p><strong>Thème:</strong> {theme}</p>
                        <p>Vous trouverez l'article complet en pièce jointe.</p>
                        <hr>
                        <h2>Aperçu de l'article:</h2>
                        <p>{structured_content['article_content'][:300]}...</p>
                        <hr>
                        <p><em>Service de génération automatique d'articles</em></p>
                        """
                        
                        send_email(
                            to_email=email_to,
                            subject=email_subject,
                            content=email_content,
                            attachment_paths=attachments,
                            is_html=True
                        )
                        
                        # Delete the temporary file after sending
                        try:
                            os.remove(output_path)
                        except:
                            pass
                    except Exception as doc_error:
                        logger.error(f"Error creating or sending document for article {uid}: {doc_error}")
                
                # Update processed count
                processed += 1
                
                # Update batch record with progress - use without await
                db.article_batches.update_one(
                    {"id": bulk_id},
                    {
                        "$set": {
                            "processed_articles": processed,
                            "status": "processing" if processed < total else "completed"
                        },
                        "$push": {"articles": article["id"]}
                    }
                )
                
            except Exception as article_error:
                logger.error(f"Error processing article with theme '{theme}': {article_error}")
        
        # Send a completion email
        if processed > 0:
            completion_subject = f"Génération d'articles terminée - {processed}/{total} articles"
            completion_content = f"""
            <h1>Génération d'articles terminée</h1>
            <p>Nous avons généré <strong>{processed}</strong> articles sur <strong>{total}</strong> demandés.</p>
            <p>Vous avez reçu chaque article dans un email séparé avec une pièce jointe au format Word.</p>
            <hr>
            <p><em>Service de génération automatique d'articles</em></p>
            """
            
            send_email(
                to_email=email_to,
                subject=completion_subject,
                content=completion_content,
                is_html=True
            )
        
        # Mark batch as completed - use without await
        db.article_batches.update_one(
            {"id": bulk_id},
            {
                "$set": {
                    "status": "completed",
                    "completed_at": datetime.now()
                }
            }
        )
        
    except Exception as e:
        logger.error(f"Error processing bulk articles: {e}")
        # Mark batch as failed - use without await
        db.get_db().article_batches.update_one(
            {"id": bulk_id},
            {
                "$set": {
                    "status": "failed",
                    "completed_at": datetime.now(),
                    "error": str(e)
                }
            }
        )

def generate_api_key():
    """Generate a secure API key"""
    return base64.b32encode(os.urandom(30)).decode('utf-8')

@app.post("/users/generate-api-key", status_code=status.HTTP_200_OK)
async def generate_user_api_key(current_user: UserInDB = Depends(get_current_user)):
    """
    Generate a new API key for the current user
    """
    db = get_db()
    
    # Generate new API key
    new_api_key = generate_api_key()
    
    # Update the user's API key in the database
    result = db.users.update_one(
        {"email": current_user.email},
        {"$set": {"api_key": new_api_key}}
    )
    
    if result.modified_count == 0:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Failed to generate API key"
        )
    
    return {"api_key": new_api_key}

@app.delete("/users/revoke-api-key", status_code=status.HTTP_200_OK)
async def revoke_user_api_key(current_user: UserInDB = Depends(get_current_user)):
    """
    Revoke the current user's API key
    """
    db = get_db()
    
    # Remove the user's API key from the database
    result = db.users.update_one(
        {"email": current_user.email},
        {"$unset": {"api_key": ""}}
    )
    
    if result.modified_count == 0:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Failed to revoke API key"
        )
    
    return {"message": "API key revoked successfully"}

@app.get("/users/api-key", status_code=status.HTTP_200_OK)
async def get_user_api_key(current_user: UserInDB = Depends(get_current_user)):
    db = get_db()
    user = db.users.find_one({"email": current_user.email})
    if not user or "api_key" not in user:
        return {"api_key": None}
    return {"api_key": user["api_key"]}

# Endpoint to set user's OpenAI API key
@app.post("/users/set-openai-api-key", status_code=status.HTTP_200_OK)
async def set_user_openai_api_key(
    payload: dict = Body(...), 
    current_user: UserInDB = Depends(get_current_user)
):
    api_key = payload.get("api_key")
    if not api_key:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="API key is required"
        )
    
    db = get_db()
    result = db.users.update_one(
        {"email": current_user.email},
        {"$set": {"openai_api_key": api_key}}
    )
    if result.modified_count == 0:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Failed to set OpenAI API key"
        )
    return {"message": "OpenAI API key set successfully"}

# Endpoint to get user's OpenAI API key
@app.get("/users/openai-api-key", status_code=status.HTTP_200_OK)
async def get_user_openai_api_key(current_user: UserInDB = Depends(get_current_user)):
    db = get_db()
    user = db.users.find_one({"email": current_user.email})
    if not user or "openai_api_key" not in user:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="No OpenAI API key found"
        )
    return {"openai_api_key": user["openai_api_key"]}